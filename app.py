"""
=============================================================
 PDF CHAPTER SPLITTER & TOOLS  —  Flask Backend  (app.py)
=============================================================
 Author  : PDF Splitter Project
 Purpose : All PDF operations are handled server-side using
           PyMuPDF (fitz). No external websites are called.

 Routes:
   /                  → Main HTML page
   /split             → Split PDF by bookmarks/chapters
   /bookmarks         → Read PDF bookmarks list
   /extract-range     → Extract specific page range
   /preview-range     → Stream page slice for iframe preview
   /preview-full      → Stream full PDF for iframe
   /download/<name>   → Download a generated file

 Tool Routes (all under /tool/):
   merge, split, compress, rotate, protect, unlock,
   watermark, page-numbers, remove-pages, reorder,
   pdf-to-images, images-to-pdf, img-to-pdf,
   html-to-pdf, website-to-pdf, pdf-info,
   convert-file

 Requirements:  flask  pymupdf  pillow  python-docx  openpyxl
=============================================================
"""

from flask import Flask, render_template, request, send_file, jsonify, Response
import fitz                          # PyMuPDF — core PDF engine
import os, re, zipfile, io, csv, json, tempfile, subprocess
from PIL import Image                # Pillow — image format conversions
from docx import Document            # python-docx — Word output
import openpyxl                      # Excel output

app = Flask(__name__)

# ── Storage folders (auto-created on startup) ──
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


# ─────────────────────────────────────────────
#  UTILITY FUNCTIONS
# ─────────────────────────────────────────────

def safe_filename(title):
    """Strip characters that are illegal in filenames."""
    s = re.sub(r'[\\/*?:"<>|]', "", str(title))
    return s.replace(" ", "_").strip("_") or "file"


def clean_text(text):
    """Remove control characters that Excel/Word cannot handle."""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)


def out_path(name):
    """Return full path inside outputs/ folder."""
    return os.path.join(OUTPUT_FOLDER, name)


def up_path(name):
    """Return full path inside uploads/ folder."""
    return os.path.join(UPLOAD_FOLDER, name)


def save_upload(file, prefix=""):
    """Save an uploaded file and return its path."""
    path = up_path(prefix + file.filename)
    file.save(path)
    return path


def stream_pdf(buf, filename):
    """
    Send a BytesIO PDF buffer as an attachment download.
    Works for both BytesIO objects and file paths.
    """
    if isinstance(buf, (str, bytes, os.PathLike)):
        return send_file(buf, as_attachment=True, download_name=filename)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf"
    )


# ─────────────────────────────────────────────
#  CHAPTER / BOOKMARK DETECTION
# ─────────────────────────────────────────────

def detect_units(doc):
    """
    Read Table of Contents (TOC) bookmarks from a PDF.
    Falls back to scanning page text for UNIT / CHAPTER keywords
    if no TOC is found (e.g. older scanned textbooks).
    Returns list of (title, 0-based-page-index) tuples sorted by page.
    """
    units = []
    toc = doc.get_toc()

    if toc:
        # Use level-1 bookmarks only (top-level chapters)
        for level, title, page in toc:
            if level == 1:
                units.append((title, page - 1))   # convert to 0-based
    else:
        # Fallback: keyword scan
        for i, page in enumerate(doc):
            text = page.get_text().upper()
            if "UNIT" in text or "CHAPTER" in text:
                lines = text.split("\n")
                for j, line in enumerate(lines):
                    if ("UNIT" in line or "CHAPTER" in line) and j + 1 < len(lines):
                        t = lines[j + 1].strip()
                        if t:
                            units.append((t, i))

    # Deduplicate while preserving order
    seen = set()
    units = [(t, p) for t, p in units if not (t in seen or seen.add(t))]
    units.sort(key=lambda x: x[1])
    return units


def generate_unit_pdfs(pdf_path, course_name):
    """
    Split a PDF into one file per chapter using bookmark data.
    Saves each chapter PDF to outputs/ folder.
    Returns list of chapter info dicts.
    """
    doc   = fitz.open(pdf_path)
    units = detect_units(doc)
    total = len(doc)
    results = []

    for i, (title, start) in enumerate(units):
        # End page = start of next chapter minus 1 (or last page)
        end = units[i + 1][1] - 1 if i < len(units) - 1 else total - 1
        end = min(end, total - 1)

        new_pdf = fitz.open()
        new_pdf.insert_pdf(doc, from_page=start, to_page=end)

        filename = f"{str(i + 1).zfill(2)}_{safe_filename(title)}.pdf"
        new_pdf.save(out_path(filename))
        new_pdf.close()

        results.append({
            "index":      i + 1,
            "title":      title,
            "filename":   filename,
            "start_page": start + 1,          # 1-based for display
            "end_page":   end + 1,
            "page_count": end - start + 1,
        })

    doc.close()
    return results


# ─────────────────────────────────────────────
#  MAIN ROUTES
# ─────────────────────────────────────────────

@app.route("/")
def index():
    """Serve the main single-page application."""
    return render_template("index.html")


@app.route("/split", methods=["POST"])
def split():
    """
    Upload a PDF → detect chapters from TOC → split into individual PDFs
    → build a ZIP → return JSON with chapter list + ZIP filename.
    """
    if "pdf" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["pdf"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Please upload a PDF file"}), 400

    # Create a clean course name from the filename
    course = (
        re.sub(r'[^a-zA-Z0-9_\- ]', '', file.filename.rsplit(".", 1)[0])
        .strip()
        .replace(" ", "_")
    ) or "course"

    path = save_upload(file)

    try:
        units = generate_unit_pdfs(path, course)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    if not units:
        return jsonify({"error": "No chapters or bookmarks found in this PDF."}), 400

    # Pack all chapter PDFs into a single ZIP for bulk download
    zip_name = f"{course}_All_Chapters.zip"
    with zipfile.ZipFile(out_path(zip_name), "w") as zf:
        for u in units:
            fp = out_path(u["filename"])
            if os.path.exists(fp):
                zf.write(fp, u["filename"])

    return jsonify({
        "course_name":  course,
        "units":        units,
        "zip_filename": zip_name,
        "total_pages":  sum(u["page_count"] for u in units),
        "saved_path":   path,
    })


@app.route("/bookmarks", methods=["POST"])
def bookmarks():
    """
    Upload a PDF → return its full TOC bookmark list as JSON.
    Used by Tab 2 (Bookmark Extract) to populate the left panel.
    """
    if "pdf" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["pdf"]
    path = save_upload(file, prefix="bk_")

    doc   = fitz.open(path)
    toc   = doc.get_toc()
    total = len(doc)
    doc.close()

    return jsonify({
        "bookmarks":  [{"level": lv, "title": t, "page": pg} for lv, t, pg in toc],
        "total_pages": total,
        "filename":   file.filename,
        "saved_path": path,
    })


@app.route("/extract-range", methods=["POST"])
def extract_range():
    """
    Extract a page range from a previously uploaded PDF.
    Body JSON: { saved_path, start_page, end_page, title }
    Returns: { filename } — the saved output file name.
    """
    data       = request.json
    path       = data.get("saved_path", "")
    start_page = int(data.get("start_page", 1))
    end_page   = int(data.get("end_page",   1))
    title      = data.get("title", "extract")

    if not os.path.exists(path):
        return jsonify({"error": "PDF not found on server. Please re-upload."}), 400

    doc   = fitz.open(path)
    total = len(doc)
    s     = max(0, start_page - 1)
    e     = min(total - 1, end_page - 1)

    new_pdf = fitz.open()
    new_pdf.insert_pdf(doc, from_page=s, to_page=e)

    filename = safe_filename(title) + ".pdf"
    new_pdf.save(out_path(filename))
    new_pdf.close()
    doc.close()

    return jsonify({"filename": filename})


@app.route("/preview-range")
def preview_range():
    """
    Stream a page-range slice of a PDF directly into an iframe.
    Query params: path=<saved_path>  start=<1-based>  end=<1-based>
    """
    path = request.args.get("path", "")
    if not path or not os.path.exists(path):
        return "PDF not found", 404

    s = max(0, int(request.args.get("start", 1)) - 1)
    e_arg = int(request.args.get("end", 1)) - 1

    doc = fitz.open(path)
    e   = min(doc.page_count - 1, e_arg)

    new_pdf = fitz.open()
    new_pdf.insert_pdf(doc, from_page=s, to_page=e)

    buf = io.BytesIO()
    new_pdf.save(buf)
    new_pdf.close()
    doc.close()
    buf.seek(0)

    return Response(
        buf.read(),
        mimetype="application/pdf",
        headers={"Content-Disposition": "inline"}
    )


@app.route("/preview-full")
def preview_full():
    """Serve a complete uploaded PDF inline for the iframe."""
    path = request.args.get("path", "")
    if not path or not os.path.exists(path):
        return "PDF not found", 404
    return send_file(path, mimetype="application/pdf")


@app.route("/download/<filename>")
def download(filename):
    """Download any generated file from the outputs/ folder."""
    fp = out_path(filename)
    if not os.path.exists(fp):
        return "File not found", 404
    return send_file(fp, as_attachment=True)


# ═════════════════════════════════════════════
#  PDF TOOLS  (all local, no external requests)
# ═════════════════════════════════════════════

# ── TOOL 1: MERGE ──────────────────────────────────────
@app.route("/tool/merge", methods=["POST"])
def tool_merge():
    """
    Merge multiple PDF files into a single PDF.
    Input:  files[] named 'pdfs'  (minimum 2)
    Output: merged.pdf
    """
    files = request.files.getlist("pdfs")
    if len(files) < 2:
        return jsonify({"error": "Upload at least 2 PDF files"}), 400

    merged = fitz.open()
    for f in files:
        buf = io.BytesIO(f.read())
        doc = fitz.open(stream=buf, filetype="pdf")
        merged.insert_pdf(doc)
        doc.close()

    buf = io.BytesIO()
    merged.save(buf)
    merged.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="merged.pdf", mimetype="application/pdf")


# ── TOOL 2: SPLIT ──────────────────────────────────────
@app.route("/tool/split", methods=["POST"])
def tool_split():
    """
    Split a PDF into separate files.
    mode='pages'  → every page as individual PDF  (ZIP output)
    mode='range'  → custom ranges like '1-5,6-10' (ZIP output)
    Input:  pdf file, mode, ranges (for range mode)
    Output: split_pages.zip
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    mode  = request.form.get("mode", "pages")
    doc   = fitz.open(stream=file.read(), filetype="pdf")
    total = doc.page_count

    zb = io.BytesIO()
    with zipfile.ZipFile(zb, "w") as zf:

        if mode == "range":
            # Parse custom range string e.g. "1-5, 6-10, 11"
            ranges_str = request.form.get("ranges", "")
            for part in ranges_str.split(","):
                part = part.strip()
                if "-" in part:
                    a, b = part.split("-", 1)
                    a, b = int(a.strip()) - 1, int(b.strip()) - 1
                elif part.isdigit():
                    a = b = int(part) - 1
                else:
                    continue

                a = max(0, a)
                b = min(total - 1, b)

                nd  = fitz.open()
                nd.insert_pdf(doc, from_page=a, to_page=b)
                pb  = io.BytesIO()
                nd.save(pb)
                nd.close()
                zf.writestr(f"pages_{a+1}_to_{b+1}.pdf", pb.getvalue())
        else:
            # Every page as a separate file
            for i in range(total):
                nd = fitz.open()
                nd.insert_pdf(doc, from_page=i, to_page=i)
                pb = io.BytesIO()
                nd.save(pb)
                nd.close()
                zf.writestr(f"page_{i+1:04d}.pdf", pb.getvalue())

    doc.close()
    zb.seek(0)
    return send_file(zb, as_attachment=True,
                     download_name="split_pages.zip", mimetype="application/zip")


# ── TOOL 3: COMPRESS ───────────────────────────────────
@app.route("/tool/compress", methods=["POST"])
def tool_compress():
    """
    Reduce PDF file size using garbage collection and deflate compression.
    Best results on text-heavy PDFs; image-heavy PDFs may not shrink much.
    Input:  pdf file
    Output: compressed.pdf
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    doc = fitz.open(stream=file.read(), filetype="pdf")
    buf = io.BytesIO()
    # garbage=4  removes unused objects; deflate=True compresses streams; clean=True sanitises
    doc.save(buf, garbage=4, deflate=True, clean=True)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="compressed.pdf", mimetype="application/pdf")


# ── TOOL 4: ROTATE ─────────────────────────────────────
@app.route("/tool/rotate", methods=["POST"])
def tool_rotate():
    """
    Rotate pages of a PDF by 90 / 180 / 270 degrees.
    Input:  pdf file, angle (90/180/270), pages ('all' or '1,3,5')
    Output: rotated.pdf
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    angle     = int(request.form.get("angle", 90))
    pages_str = request.form.get("pages", "all").strip()

    doc = fitz.open(stream=file.read(), filetype="pdf")
    for i, page in enumerate(doc):
        # Rotate all pages OR only specified page numbers
        if pages_str == "all" or str(i + 1) in [p.strip() for p in pages_str.split(",")]:
            page.set_rotation((page.rotation + angle) % 360)

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="rotated.pdf", mimetype="application/pdf")


# ── TOOL 5: PROTECT ────────────────────────────────────
@app.route("/tool/protect", methods=["POST"])
def tool_protect():
    """
    Add AES-256 password protection to a PDF.
    Input:  pdf file, password
    Output: protected.pdf
    """
    file     = request.files.get("pdf")
    password = request.form.get("password", "")

    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400
    if not password:
        return jsonify({"error": "Password is required"}), 400

    doc  = fitz.open(stream=file.read(), filetype="pdf")
    perm = fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY    # allow print + copy
    buf  = io.BytesIO()
    doc.save(
        buf,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw=password + "_owner",
        user_pw=password,
        permissions=perm
    )
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="protected.pdf", mimetype="application/pdf")


# ── TOOL 6: UNLOCK ─────────────────────────────────────
@app.route("/tool/unlock", methods=["POST"])
def tool_unlock():
    """
    Remove password protection from a PDF (requires current password).
    Input:  pdf file, password
    Output: unlocked.pdf
    """
    file     = request.files.get("pdf")
    password = request.form.get("password", "")

    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    raw = file.read()
    doc = fitz.open(stream=raw, filetype="pdf")

    if doc.needs_pass:
        if not doc.authenticate(password):
            return jsonify({"error": "Incorrect password"}), 400

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="unlocked.pdf", mimetype="application/pdf")


# ── TOOL 7: WATERMARK ──────────────────────────────────
@app.route("/tool/watermark", methods=["POST"])
def tool_watermark():
    """
    Add a diagonal text watermark to every page.
    Input:  pdf, text, color (red/blue/gray/black/green), opacity (0.1-1.0), size
    Output: watermarked.pdf
    """
    file    = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    text    = request.form.get("text", "WATERMARK")
    color   = request.form.get("color", "gray")
    opacity = float(request.form.get("opacity", "0.3"))
    size    = int(request.form.get("size", "60"))

    # Named colour → RGB tuple
    colour_map = {
        "red":   (1, 0, 0),
        "blue":  (0, 0, 1),
        "gray":  (0.5, 0.5, 0.5),
        "black": (0, 0, 0),
        "green": (0, 0.5, 0),
    }
    rgb = colour_map.get(color, (0.5, 0.5, 0.5))

    doc = fitz.open(stream=file.read(), filetype="pdf")
    for page in doc:
        r    = page.rect
        tw   = fitz.TextWriter(r)
        font = fitz.Font("helv")
        # Centre the watermark text approximately
        x = r.width  / 2 - len(text) * size * 0.28
        y = r.height / 2
        tw.append((x, y), text, font=font, fontsize=size)
        # Rotate 45° around page centre
        tw.write_text(
            page,
            color=rgb,
            opacity=opacity,
            morph=(r.center, fitz.Matrix(-45))
        )

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="watermarked.pdf", mimetype="application/pdf")


# ── TOOL 8: PAGE NUMBERS ───────────────────────────────
@app.route("/tool/page-numbers", methods=["POST"])
def tool_page_numbers():
    """
    Add automatic page numbers to every page.
    Input:  pdf, position, start_from, prefix
    Output: numbered.pdf
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    pos    = request.form.get("position",   "bottom-center")
    start  = int(request.form.get("start_from", 1))
    prefix = request.form.get("prefix", "")

    doc = fitz.open(stream=file.read(), filetype="pdf")
    font = fitz.Font("helv")
    fs   = 11    # font size

    for i, page in enumerate(doc):
        label  = f"{prefix}{start + i}"
        r      = page.rect
        # Estimate text width for centering
        text_w = len(label) * fs * 0.5

        # Map position name → (x, y) coordinate
        position_map = {
            "bottom-center": (r.width / 2 - text_w / 2, r.height - 20),
            "bottom-right":  (r.width - text_w - 20,     r.height - 20),
            "bottom-left":   (20,                          r.height - 20),
            "top-center":    (r.width / 2 - text_w / 2, 25),
            "top-right":     (r.width - text_w - 20,     25),
            "top-left":      (20,                          25),
        }
        pt = position_map.get(pos, position_map["bottom-center"])

        tw = fitz.TextWriter(r)
        tw.append(pt, label, font=font, fontsize=fs)
        tw.write_text(page, color=(0.3, 0.3, 0.3))

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="numbered.pdf", mimetype="application/pdf")


# ── TOOL 9: REMOVE PAGES ───────────────────────────────
@app.route("/tool/remove-pages", methods=["POST"])
def tool_remove_pages():
    """
    Delete specified pages from a PDF.
    Input:  pdf, pages (e.g. '1, 3, 5-8')
    Output: pages_removed.pdf
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    pages_str = request.form.get("pages", "")
    doc       = fitz.open(stream=file.read(), filetype="pdf")
    to_remove = set()

    for part in pages_str.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-")
            to_remove.update(range(int(a) - 1, int(b)))
        elif part.isdigit():
            to_remove.add(int(part) - 1)

    # Delete in reverse order so indices remain valid
    for p in sorted(to_remove, reverse=True):
        if 0 <= p < doc.page_count:
            doc.delete_page(p)

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="pages_removed.pdf", mimetype="application/pdf")


# ── TOOL 10: REORDER PAGES ─────────────────────────────
@app.route("/tool/reorder", methods=["POST"])
def tool_reorder():
    """
    Rearrange pages in a custom order.
    Input:  pdf, order (e.g. '3,1,2' for 3-page PDF)
    Output: reordered.pdf
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    order_str = request.form.get("order", "")
    doc       = fitz.open(stream=file.read(), filetype="pdf")

    try:
        order = [int(x.strip()) - 1 for x in order_str.split(",") if x.strip().isdigit()]
        order = [o for o in order if 0 <= o < doc.page_count]
    except Exception:
        return jsonify({"error": "Invalid page order format"}), 400

    doc.select(order)
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="reordered.pdf", mimetype="application/pdf")


# ── TOOL 11: PDF → IMAGES ──────────────────────────────
@app.route("/tool/pdf-to-images", methods=["POST"])
def tool_pdf_to_images():
    """
    Convert every PDF page to an image file (JPG or PNG).
    Input:  pdf, format ('jpg'/'png'), dpi (72/150/200/300)
    Output: pdf_images.zip
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    fmt = request.form.get("format", "jpg").lower()
    dpi = int(request.form.get("dpi", "150"))
    mat = fitz.Matrix(dpi / 72, dpi / 72)   # scale factor

    doc = fitz.open(stream=file.read(), filetype="pdf")
    zb  = io.BytesIO()
    ext = "jpg" if fmt == "jpg" else fmt

    with zipfile.ZipFile(zb, "w") as zf:
        for i, page in enumerate(doc):
            pix      = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("jpeg" if fmt == "jpg" else fmt)
            zf.writestr(f"page_{i+1:04d}.{ext}", img_data)

    doc.close()
    zb.seek(0)
    return send_file(zb, as_attachment=True,
                     download_name="pdf_images.zip", mimetype="application/zip")


# ── TOOL 12: IMAGES → PDF ──────────────────────────────
@app.route("/tool/images-to-pdf", methods=["POST"])
def tool_images_to_pdf():
    """
    Combine multiple image files into a single PDF.
    Input:  images[] (JPG / PNG)
    Output: images_to_pdf.pdf
    """
    files = request.files.getlist("images")
    if not files:
        return jsonify({"error": "No images uploaded"}), 400

    doc = fitz.open()
    for f in files:
        img_bytes = f.read()
        ext       = f.filename.rsplit(".", 1)[-1].lower()
        img_doc   = fitz.open(stream=img_bytes, filetype=ext)
        pdfbytes  = img_doc.convert_to_pdf()
        img_doc.close()
        imgpdf = fitz.open("pdf", pdfbytes)
        doc.insert_pdf(imgpdf)
        imgpdf.close()

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="images_to_pdf.pdf", mimetype="application/pdf")


# ── TOOL 13: SINGLE IMAGE → PDF ────────────────────────
@app.route("/tool/img-to-pdf", methods=["POST"])
def tool_img_to_pdf():
    """
    Convert a single JPG or PNG image to PDF.
    Input:  image file
    Output: converted.pdf
    """
    file = request.files.get("image")
    if not file:
        return jsonify({"error": "No image file uploaded"}), 400

    img_bytes = file.read()
    ext       = file.filename.rsplit(".", 1)[-1].lower()
    img_doc   = fitz.open(stream=img_bytes, filetype=ext)
    pdfbytes  = img_doc.convert_to_pdf()
    img_doc.close()

    imgpdf = fitz.open("pdf", pdfbytes)
    buf    = io.BytesIO()
    imgpdf.save(buf)
    imgpdf.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="converted.pdf", mimetype="application/pdf")


# ── TOOL 14: HTML → PDF ────────────────────────────────
@app.route("/tool/html-to-pdf", methods=["POST"])
def tool_html_to_pdf():
    """
    Convert an HTML string to a PDF page using fitz.insert_htmlbox.
    Input:  html (form field with HTML content)
    Output: from_html.pdf
    """
    html = (request.form.get("html") or "").strip()
    if not html:
        return jsonify({"error": "No HTML content provided"}), 400

    doc  = fitz.open()
    page = doc.new_page(width=595, height=842)   # A4 size
    rect = fitz.Rect(40, 40, 555, 802)           # margins
    page.insert_htmlbox(rect, html, css="body{font-family:sans-serif;font-size:12pt}")

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="from_html.pdf", mimetype="application/pdf")


# ── TOOL 15: WEBSITE URL → PDF ─────────────────────────
@app.route("/tool/website-to-pdf", methods=["POST"])
def tool_website_to_pdf():
    """
    Convert a webpage to PDF using wkhtmltopdf (installed system tool).
    The browser fetches the page HTML and sends it here to avoid
    network restrictions in server environments.
    Input:  html (full page HTML sent from browser), url (for filename)
    Output: website.pdf
    """
    html = (request.form.get("html") or "").strip()
    url  = (request.form.get("url")  or "website").strip()

    if not html:
        return jsonify({"error": "No page content received. Try a different website."}), 400

    # Write HTML to a temp file, then run wkhtmltopdf on it
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as f:
        f.write(html)
        html_path = f.name

    pdf_path = html_path.replace(".html", ".pdf")

    try:
        result = subprocess.run(
            ["wkhtmltopdf", "--quiet", "--enable-local-file-access",
             html_path, pdf_path],
            capture_output=True, text=True, timeout=60
        )

        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            # Fallback: use fitz insert_htmlbox
            doc  = fitz.open()
            page = doc.new_page(width=595, height=842)
            rect = fitz.Rect(36, 36, 559, 806)
            page.insert_htmlbox(rect, html[:50000])   # cap at 50k chars
            doc.save(pdf_path)
            doc.close()

    finally:
        # Clean up temp HTML file
        try:
            os.unlink(html_path)
        except Exception:
            pass

    safe_url = re.sub(r'[^a-zA-Z0-9_\-]', '_', url)[:40]
    return send_file(
        pdf_path,
        as_attachment=True,
        download_name=f"{safe_url}.pdf",
        mimetype="application/pdf"
    )


# ── TOOL 16: FILE CONVERTER ────────────────────────────
@app.route("/tool/convert-file", methods=["POST"])
def tool_convert_file():
    """
    Universal file converter.  Upload ANY supported file and choose
    an output format from the dropdown.

    Supported INPUT  → OUTPUT combinations:
      PDF        → Word (.docx), Excel (.xlsx), CSV, PNG, JPG, TXT
      Image      → PDF, PNG, JPG, WEBP, BMP, TIFF
      CSV        → PDF, Excel (.xlsx)
      TXT        → PDF
      Word       → PDF (via HTML conversion)

    Input:  file, output_format (dropdown selection)
    Output: converted file download
    """
    file          = request.files.get("file")
    output_format = request.form.get("output_format", "pdf").lower()

    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    filename = file.filename
    ext      = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    raw      = file.read()

    # ── PDF  →  anything ──
    if ext == "pdf":
        doc = fitz.open(stream=raw, filetype="pdf")

        if output_format in ("word", "docx"):
            # Extract text from all pages → Word document
            wdoc = Document()
            wdoc.add_heading("Extracted from " + filename, level=1)
            for i, page in enumerate(doc):
                wdoc.add_heading(f"Page {i + 1}", level=2)
                wdoc.add_paragraph(clean_text(page.get_text("text")))
            doc.close()
            buf = io.BytesIO()
            wdoc.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.docx",
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif output_format in ("excel", "xlsx"):
            # Extract text per page → Excel rows
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PDF Content"
            ws.append(["Page Number", "Content"])
            for i, page in enumerate(doc):
                text = clean_text(page.get_text("text").replace("\n", " "))[:500]
                ws.append([i + 1, text])
            doc.close()
            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.xlsx",
                             mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        elif output_format == "csv":
            # Extract text per page → CSV
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(["Page", "Content"])
            for i, page in enumerate(doc):
                text = clean_text(page.get_text("text").replace("\n", " "))[:500]
                writer.writerow([i + 1, text])
            doc.close()
            buf = io.BytesIO(output.getvalue().encode("utf-8"))
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.csv", mimetype="text/csv")

        elif output_format == "txt":
            # Extract all text → plain text file
            texts = []
            for i, page in enumerate(doc):
                texts.append(f"=== Page {i+1} ===\n{clean_text(page.get_text('text'))}")
            doc.close()
            buf = io.BytesIO("\n\n".join(texts).encode("utf-8"))
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.txt", mimetype="text/plain")

        elif output_format in ("png", "jpg", "jpeg"):
            # Convert first page to image
            pix      = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
            fmt      = "jpeg" if output_format in ("jpg", "jpeg") else "png"
            img_data = pix.tobytes(fmt)
            doc.close()
            buf = io.BytesIO(img_data)
            buf.seek(0)
            mime = "image/jpeg" if fmt == "jpeg" else "image/png"
            return send_file(buf, as_attachment=True,
                             download_name=f"page_1.{output_format}", mimetype=mime)

        doc.close()
        return jsonify({"error": f"PDF → {output_format} not supported"}), 400

    # ── IMAGE  →  anything ──
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif", "gif"):
        pil_img = Image.open(io.BytesIO(raw)).convert("RGBA")

        if output_format == "pdf":
            # Image → PDF via fitz
            img_doc  = fitz.open(stream=raw, filetype=ext if ext not in ("jpg",) else "jpeg")
            pdfbytes = img_doc.convert_to_pdf()
            img_doc.close()
            buf = io.BytesIO(pdfbytes)
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.pdf", mimetype="application/pdf")

        fmt_map = {
            "png":  ("PNG",  "image/png"),
            "jpg":  ("JPEG", "image/jpeg"),
            "jpeg": ("JPEG", "image/jpeg"),
            "webp": ("WEBP", "image/webp"),
            "bmp":  ("BMP",  "image/bmp"),
            "tiff": ("TIFF", "image/tiff"),
        }
        if output_format in fmt_map:
            pil_fmt, mime = fmt_map[output_format]
            img_out = pil_img.convert("RGB") if pil_fmt == "JPEG" else pil_img
            buf = io.BytesIO()
            img_out.save(buf, pil_fmt)
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name=f"converted.{output_format}", mimetype=mime)

        return jsonify({"error": f"Image → {output_format} not supported"}), 400

    # ── CSV  →  anything ──
    elif ext == "csv":
        text_content = raw.decode("utf-8", errors="replace")
        reader       = csv.reader(io.StringIO(text_content))
        rows         = list(reader)

        if output_format == "pdf":
            # Build HTML table → PDF
            table_html = '<table border="1" cellpadding="5" style="border-collapse:collapse;font-size:10pt">'
            for i, row in enumerate(rows):
                tag = "th" if i == 0 else "td"
                table_html += "<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in row) + "</tr>"
            table_html += "</table>"
            doc  = fitz.open()
            page = doc.new_page(width=842, height=595)   # landscape A4
            page.insert_htmlbox(fitz.Rect(20, 20, 822, 575), table_html)
            buf = io.BytesIO()
            doc.save(buf)
            doc.close()
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.pdf", mimetype="application/pdf")

        elif output_format in ("excel", "xlsx"):
            wb = openpyxl.Workbook()
            ws = wb.active
            for row in rows:
                ws.append([clean_text(c) for c in row])
            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.xlsx",
                             mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        return jsonify({"error": f"CSV → {output_format} not supported"}), 400

    # ── TXT  →  PDF ──
    elif ext == "txt":
        text = raw.decode("utf-8", errors="replace")
        if output_format == "pdf":
            doc  = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_textbox(
                fitz.Rect(40, 40, 555, 802),
                clean_text(text),
                fontname="helv",
                fontsize=11
            )
            buf = io.BytesIO()
            doc.save(buf)
            doc.close()
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.pdf", mimetype="application/pdf")
        return jsonify({"error": f"TXT → {output_format} not supported"}), 400

    # ── WORD (.docx) → PDF ──
    elif ext == "docx":
        if output_format == "pdf":
            wdoc     = Document(io.BytesIO(raw))
            paragraphs = "\n".join(p.text for p in wdoc.paragraphs)
            doc  = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_textbox(
                fitz.Rect(40, 40, 555, 802),
                clean_text(paragraphs),
                fontname="helv",
                fontsize=11
            )
            buf = io.BytesIO()
            doc.save(buf)
            doc.close()
            buf.seek(0)
            return send_file(buf, as_attachment=True,
                             download_name="converted.pdf", mimetype="application/pdf")
        return jsonify({"error": f"Word → {output_format} not supported"}), 400

    return jsonify({"error": f"Unsupported input file type: .{ext}"}), 400


# ── TOOL 17: PDF INFO ──────────────────────────────────
@app.route("/tool/pdf-info", methods=["POST"])
def tool_pdf_info():
    """
    Return metadata and statistics for a PDF file as JSON.
    Input:  pdf file
    Output: JSON with pages, title, author, size, encryption, etc.
    """
    file = request.files.get("pdf")
    if not file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    raw  = file.read()
    doc  = fitz.open(stream=raw, filetype="pdf")
    meta = doc.metadata or {}
    toc  = doc.get_toc()

    info = {
        "pages":     doc.page_count,
        "title":     meta.get("title",    "—"),
        "author":    meta.get("author",   "—"),
        "subject":   meta.get("subject",  "—"),
        "creator":   meta.get("creator",  "—"),
        "producer":  meta.get("producer", "—"),
        "encrypted": doc.needs_pass,
        "bookmarks": len(toc),
        "size_kb":   round(len(raw) / 1024, 1),
        "page_size": (
            f"{doc[0].rect.width:.0f} × {doc[0].rect.height:.0f} pt"
            if doc.page_count else "—"
        ),
    }
    doc.close()
    return jsonify(info)


# ─────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "=" * 54)
    print("  PDF Chapter Splitter & Tools — Running!")
    print("  Open your browser at: http://127.0.0.1:5000")
    print("=" * 54 + "\n")
    app.run(debug=True, port=5000)
